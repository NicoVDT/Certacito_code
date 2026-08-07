#!/usr/bin/env python3
"""
Build the A4 appendix set as real .docx files.

The submission portal is fussy about pdf uploads, so these are genuine Word
documents rather than a pdf with its metadata rewritten - same content, same
branding as build-pdfs.mjs, and Word can re-export them to pdf if needed.
"""
import os
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(DOCS_DIR, "docx")

NAVY = RGBColor(0x1B, 0x3A, 0x6B)
TEAL = RGBColor(0x0D, 0x73, 0x77)
INK = RGBColor(0x1A, 0x1A, 0x2E)
GREY = RGBColor(0x6B, 0x7A, 0x99)

DOCUMENTS = [
    ("requirements-traceability.md", "Requirements Traceability Matrix", True),
    ("sprint-progress.md", "Sprint Progress Against the A2 Plan", True),
    ("architecture.md", "System Architecture", False),
    ("functionality.md", "System Functionality", False),
    ("interface-design.md", "Interface Design", False),
    ("usability-evaluation.md", "Usability Evaluation", False),
    ("branding-style-guide.md", "Branding and Style Guide", False),
    ("contribution-table.md", "Contribution Table", False),
]


def shade(cell, hex_colour):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(el)


def add_runs(para, text):
    """Inline markdown: **bold**, `code`, *italic*. Links become their label."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
            r.font.color.rgb = NAVY
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(9)
            r.font.color.rgb = NAVY
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            para.add_run(part[1:-1]).italic = True
        else:
            para.add_run(part)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def cover(doc, title, source_name):
    for _ in range(6):
        doc.add_paragraph()

    logo = os.path.join(DOCS_DIR, "logo-print.png")
    if os.path.exists(logo):
        doc.add_picture(logo, width=Inches(0.85))

    p = doc.add_paragraph()
    r = p.add_run("GROUP 28  ·  CSIT321")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = TEAL

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run("Certacito.ai  ·  A4 Prototype Presentation appendix")
    r.font.size = Pt(12)
    r.font.color.rgb = GREY

    doc.add_paragraph()
    for label, value in [
        ("Project", "Certacito.ai, AI agent governance platform"),
        ("Document", source_name),
        ("Date", "07 August 2026"),
        ("Live system", "http://20.92.93.30"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}   ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
        r = p.add_run(value)
        r.font.size = Pt(10)
        r.font.color.rgb = GREY

    doc.add_page_break()


def footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CERTACITO.AI          University of Wollongong | CSIT321 | 2026 | Group 28")
    r.font.size = Pt(7.5)
    r.font.color.rgb = GREY


def build(md_name, title, landscape):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK

    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ("top_margin", "bottom_margin"):
        setattr(section, attr, Inches(0.7))
    section.left_margin = section.right_margin = Inches(0.55)
    footer(section)

    cover(doc, title, md_name)

    lines = open(os.path.join(DOCS_DIR, md_name), encoding="utf8").read().split("\n")
    lines = lines[1:] if lines and lines[0].startswith("# ") else lines

    i = 0
    in_code = []
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(in_code))
                r.font.name = "Courier New"
                r.font.size = Pt(8)
                shade_p = OxmlElement("w:shd")
                shade_p.set(qn("w:val"), "clear")
                shade_p.set(qn("w:fill"), "F4F6F9")
                p._p.get_or_add_pPr().append(shade_p)
                in_code = []
            else:
                in_code = [""]
                in_code.pop()
                in_code.append("")
                in_code.pop()
                in_code = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    in_code.append(lines[i])
                    i += 1
                continue
            i += 1
            continue

        # table
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1]):
            header = split_row(line)
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for idx, text in enumerate(header):
                cell = t.rows[0].cells[idx]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(text.upper())
                r.bold = True
                r.font.size = Pt(7.5)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade(cell, "1B3A6B")
            for n, row in enumerate(rows):
                cells = t.add_row().cells
                for idx, text in enumerate(row[: len(header)]):
                    cells[idx].text = ""
                    p = cells[idx].paragraphs[0]
                    p.paragraph_format.space_after = Pt(1)
                    add_runs(p, text)
                    for r in p.runs:
                        r.font.size = Pt(8)
                    if n % 2:
                        shade(cells[idx], "F7F9FC")
            doc.add_paragraph()
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:])
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = NAVY
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = NAVY
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, line[2:])
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x40, 0x50, 0x6B)
        elif re.match(r"^\s*[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, re.sub(r"^\s*[-*] ", "", line))
        elif re.match(r"^\s*\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, re.sub(r"^\s*\d+\. ", "", line))
        elif line.startswith("!["):
            m = re.search(r"\(([^)]+)\)", line)
            if m:
                img = os.path.join(DOCS_DIR, m.group(1))
                if os.path.exists(img):
                    doc.add_picture(img, width=Inches(9.0 if landscape else 6.4))
        elif line.strip() == "---":
            pass
        elif line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_runs(p, line)

        i += 1

    core = doc.core_properties
    core.title = f"Certacito.ai - {title}"
    core.author = "Group 28, CSIT321, University of Wollongong"
    core.subject = "A4 Prototype Presentation appendix"
    core.category = "CSIT321 A4"

    out = os.path.join(OUT_DIR, md_name.replace(".md", ".docx"))
    doc.save(out)
    return out


if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    for md_name, title, landscape in DOCUMENTS:
        path = build(md_name, title, landscape)
        print(f"wrote {os.path.basename(path)} {os.path.getsize(path) / 1024:.0f}KB")
